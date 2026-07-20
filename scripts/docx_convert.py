"""md → docx 변환 + 재추출 텍스트 대조 검증."""
import re
import sys

import docx

from charcount import CLAIM_RE


def _clean_md(md_text: str) -> str:
    return CLAIM_RE.sub("", md_text)


def md_to_docx(md_text: str, out_path: str) -> None:
    doc = docx.Document()
    cleaned = _clean_md(md_text)

    # Split by blank lines to get paragraph blocks
    for block in cleaned.split('\n\n'):
        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue

        # Process each line in the block
        para_lines = []
        for line in lines:
            m = re.match(r"^(#+) (.*)$", line)
            if m:
                # Heading: flush accumulated paragraph first, then add heading
                if para_lines:
                    doc.add_paragraph(' '.join(para_lines))
                    para_lines = []
                doc.add_heading(m.group(2), level=min(len(m.group(1)), 4))
            else:
                # Accumulate non-heading lines
                para_lines.append(line)

        # Add accumulated paragraph at end of block
        if para_lines:
            doc.add_paragraph(' '.join(para_lines))

    doc.save(out_path)


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def verify_roundtrip(md_text: str, docx_path: str) -> bool:
    body = re.sub(r"^#+ .*$", "", _clean_md(md_text), flags=re.MULTILINE)
    extracted = "\n".join(
        p.text for p in docx.Document(docx_path).paragraphs
        if not p.style.name.startswith("Heading")
    )
    return _normalize(body) == _normalize(extracted)


def main(md_path: str, out_path: str) -> int:
    md_text = open(md_path, encoding="utf-8").read()
    md_to_docx(md_text, out_path)
    ok = verify_roundtrip(md_text, out_path)
    print("변환 완료" + ("" if ok else " — 재추출 대조 실패: 수동 확인 필요"))
    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main(sys.argv[1], sys.argv[2]))
