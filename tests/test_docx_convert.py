import sys, pathlib, tempfile
sys.path.insert(0, str(pathlib.Path(__file__).parent.parent / "scripts"))
from docx_convert import md_to_docx, verify_roundtrip

MD = """## 문항 1. 지원동기
사용자 인터뷰를 진행했습니다.<!--c:DIRECT:EXP-01-->

두 번째 문단입니다."""

def test_roundtrip():
    with tempfile.TemporaryDirectory() as d:
        out = f"{d}/o.docx"
        md_to_docx(MD, out)
        assert verify_roundtrip(MD, out)

def test_claim_comments_stripped():
    import docx
    with tempfile.TemporaryDirectory() as d:
        out = f"{d}/o.docx"
        md_to_docx(MD, out)
        full = "\n".join(p.text for p in docx.Document(out).paragraphs)
        assert "c:DIRECT" not in full
        assert "사용자 인터뷰를 진행했습니다." in full

def test_multiline_paragraph_not_split():
    import docx
    md = "## 문항 1. 지원동기\n첫 줄과\n둘째 줄이 한 문단.\n\n새 문단."
    with tempfile.TemporaryDirectory() as d:
        out = f"{d}/o.docx"
        md_to_docx(md, out)
        paras = [p.text for p in docx.Document(out).paragraphs
                 if p.text and not p.style.name.startswith("Heading")]
        assert paras == ["첫 줄과 둘째 줄이 한 문단.", "새 문단."]
